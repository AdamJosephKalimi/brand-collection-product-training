"""
Document Parser Service for extracting text and images from various file formats.
"""
import io
import uuid
import hashlib
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Inches
import base64
import pandas as pd
import openpyxl
import fitz  # PyMuPDF
from fastapi import HTTPException

from ..models.document import DocumentType

class ParserService:
    """Service for parsing documents and extracting text and images"""
    
    def __init__(self):
        pass
    
    async def parse_pdf(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse PDF document and extract text and images using PyMuPDF.
        
        Args:
            file_bytes: PDF file content as bytes
            filename: Original filename
            
        Returns:
            Dictionary containing extracted text, images, and metadata
        """
        try:
            # Open PDF with PyMuPDF (fitz)
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            
            # Initialize variables
            extracted_text = []
            page_metadata = []
            total_pages = len(doc)
            
            for page_num in range(total_pages):
                page = doc[page_num]
                
                # Extract text from page
                page_text = page.get_text()
                if page_text:
                    extracted_text.append({
                        "page": page_num + 1,
                        "text": page_text.strip()
                    })
                
                # Extract tables if present
                tables = page.find_tables()
                has_tables = False
                if tables and tables.tables:
                    has_tables = True
                    for table_idx, table in enumerate(tables.tables):
                        # Convert table to text representation
                        table_data = table.extract()
                        table_text = self._table_to_text(table_data, page_num + 1, table_idx)
                        extracted_text.append({
                            "page": page_num + 1,
                            "text": table_text,
                            "type": "table"
                        })
                
                # Collect page metadata
                rect = page.rect
                page_metadata.append({
                    "page": page_num + 1,
                    "width": rect.width,
                    "height": rect.height,
                    "has_text": bool(page_text),
                    "has_tables": has_tables
                })
            
            # Close document
            doc.close()
            
            # Combine all text
            full_text = "\n\n".join([item["text"] for item in extracted_text])
            
            return {
                "document_type": DocumentType.PDF,
                "filename": filename,
                "total_pages": total_pages,
                "extracted_text": full_text,
                "text_by_page": extracted_text,
                "page_metadata": page_metadata,
                "metadata": {
                    "parser": "PyMuPDF",
                    "total_characters": len(full_text),
                    "pages_with_text": len([p for p in page_metadata if p["has_text"]]),
                    "pages_with_tables": len([p for p in page_metadata if p["has_tables"]])
                }
            }
            
        except Exception as e:
            raise HTTPException(
                status_code=500, 
                detail=f"Failed to parse PDF '{filename}': {str(e)}"
            )
    
    def _table_to_text(self, table: List[List], page_num: int, table_idx: int) -> str:
        """
        Convert extracted table to readable text format.
        
        Args:
            table: Table data as list of lists
            page_num: Page number
            table_idx: Table index on the page
            
        Returns:
            Formatted table text
        """
        if not table:
            return ""
        
        # Filter out None values and empty rows
        clean_table = []
        for row in table:
            if row and any(cell for cell in row if cell is not None):
                clean_row = [str(cell) if cell is not None else "" for cell in row]
                clean_table.append(clean_row)
        
        if not clean_table:
            return ""
        
        # Create text representation
        table_text = f"[Table {table_idx + 1} on page {page_num}]\n"
        
        # Add header if first row looks like headers
        if clean_table:
            header_row = clean_table[0]
            table_text += " | ".join(header_row) + "\n"
            table_text += "-" * len(" | ".join(header_row)) + "\n"
            
            # Add data rows
            for row in clean_table[1:]:
                table_text += " | ".join(row) + "\n"
        
        return table_text.strip()
    
    def _extract_images_with_pymupdf(self, file_bytes: bytes, page_num: int = None) -> List[Dict[str, Any]]:
        """
        Extract images from PDF using PyMuPDF for robust image extraction.
        
        Args:
            file_bytes: PDF file content as bytes
            page_num: Specific page number to extract from (1-indexed), or None for all pages
            
        Returns:
            List of image dictionaries with metadata and base64 data
        """
        images = []
        
        try:
            # Open PDF with PyMuPDF
            pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
            
            # Determine which pages to process
            if page_num is not None:
                # Extract from specific page (convert to 0-indexed)
                pages_to_process = [page_num - 1] if page_num - 1 < len(pdf_document) else []
            else:
                # Extract from all pages
                pages_to_process = range(len(pdf_document))
            
            for page_idx in pages_to_process:
                page = pdf_document[page_idx]
                current_page_num = page_idx + 1  # Convert back to 1-indexed
                
                # Get image list from page
                image_list = page.get_images()
                
                for img_idx, img in enumerate(image_list):
                    try:
                        # Get image reference
                        xref = img[0]
                        
                        # Extract image data
                        base_image = pdf_document.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        
                        # Generate SHA256 hash for deduplication
                        sha256_hash = hashlib.sha256(image_bytes).hexdigest()
                        
                        # Determine content type
                        content_type_map = {
                            "png": "image/png",
                            "jpg": "image/jpeg",
                            "jpeg": "image/jpeg",
                            "bmp": "image/bmp",
                            "tiff": "image/tiff",
                            "webp": "image/webp"
                        }
                        content_type = content_type_map.get(image_ext.lower(), "image/png")
                        
                        # Get image rectangle (position and size)
                        img_rect = page.get_image_rects(img)[0] if page.get_image_rects(img) else fitz.Rect(0, 0, 0, 0)
                        
                        # Generate unique image ID
                        image_id = str(uuid.uuid4())
                        
                        # Create storage key for Firebase
                        storage_key = f"pdf-images/{sha256_hash}.{image_ext}"
                        
                        # Create image info (metadata only, no raw bytes)
                        image_info = {
                            "image_id": image_id,
                            "page": current_page_num,
                            "index": img_idx,
                            "bbox": [img_rect.x0, img_rect.y0, img_rect.x1, img_rect.y1],
                            "width": int(img_rect.width),
                            "height": int(img_rect.height),
                            "filename": f"page_{current_page_num}_image_{img_idx + 1}.{image_ext}",
                            "content_type": content_type,
                            "size": len(image_bytes),
                            "format": image_ext.upper(),
                            "colorspace": base_image.get("colorspace", "unknown"),
                            "sha256": sha256_hash,
                            "storage_key": storage_key,
                            "_raw_bytes": image_bytes  # Temporary, for upload processing
                        }
                        
                        images.append(image_info)
                        
                    except Exception as img_error:
                        print(f"Warning: Could not extract image {img_idx} from page {current_page_num}: {str(img_error)}")
                        continue
            
            pdf_document.close()
            
        except Exception as e:
            print(f"Warning: PyMuPDF image extraction failed: {str(e)}")
        
        return images
    
    async def parse_docx(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse DOCX document and extract text, tables, and images.
        
        Args:
            file_bytes: DOCX file content as bytes
            filename: Original filename
            
        Returns:
            Dictionary containing extracted text, images, and metadata
        """
        try:
            # Create file-like object from bytes
            docx_file = io.BytesIO(file_bytes)
            
            # Initialize results
            extracted_text = []
            extracted_images = []
            extracted_tables = []
            
            # Open DOCX with python-docx
            doc = Document(docx_file)
            
            # Extract paragraphs
            paragraph_count = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    extracted_text.append({
                        "type": "paragraph",
                        "index": paragraph_count,
                        "text": para.text.strip(),
                        "style": para.style.name if para.style else "Normal"
                    })
                    paragraph_count += 1
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                # Convert table to text
                table_text = self._docx_table_to_text(table_data, table_idx)
                extracted_tables.append({
                    "type": "table",
                    "index": table_idx,
                    "text": table_text,
                    "raw_data": table_data
                })
            
            # Extract images (inline images)
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        image_id = str(uuid.uuid4())
                        
                        # Get image format from content type
                        content_type = rel.target_part.content_type
                        if 'jpeg' in content_type:
                            image_format = 'jpg'
                        elif 'png' in content_type:
                            image_format = 'png'
                        else:
                            image_format = 'img'
                        
                        extracted_images.append({
                            "image_id": image_id,
                            "index": image_count,
                            "filename": f"docx_image_{image_count + 1}.{image_format}",
                            "content_type": content_type,
                            "size": len(image_data),
                            "data": base64.b64encode(image_data).decode('utf-8')
                        })
                        image_count += 1
                        
                    except Exception as e:
                        print(f"Warning: Could not extract image {image_count}: {str(e)}")
            
            # Combine all text content
            all_text_parts = []
            
            # Add paragraphs
            for item in extracted_text:
                all_text_parts.append(item["text"])
            
            # Add tables
            for item in extracted_tables:
                all_text_parts.append(item["text"])
            
            full_text = "\n\n".join(all_text_parts)
            
            return {
                "document_type": DocumentType.DOCX,
                "filename": filename,
                "extracted_text": full_text,
                "paragraphs": extracted_text,
                "tables": extracted_tables,
                "extracted_images": extracted_images,
                "metadata": {
                    "parser": "python-docx",
                    "total_characters": len(full_text),
                    "paragraph_count": len(extracted_text),
                    "table_count": len(extracted_tables),
                    "image_count": len(extracted_images),
                    "sections": len(doc.sections) if hasattr(doc, 'sections') else 0
                }
            }
            
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to parse DOCX '{filename}': {str(e)}"
            )
    
    def _docx_table_to_text(self, table_data: List[List[str]], table_idx: int) -> str:
        """
        Convert DOCX table data to readable text format.
        
        Args:
            table_data: Table data as list of lists
            table_idx: Table index in document
            
        Returns:
            Formatted table text
        """
        if not table_data:
            return ""
        
        # Filter out empty rows
        clean_table = []
        for row in table_data:
            if any(cell.strip() for cell in row):
                clean_table.append([cell.strip() for cell in row])
        
        if not clean_table:
            return ""
        
        # Create text representation
        table_text = f"[Table {table_idx + 1}]\n"
        
        # Add header if first row looks like headers
        if clean_table:
            header_row = clean_table[0]
            table_text += " | ".join(header_row) + "\n"
            table_text += "-" * len(" | ".join(header_row)) + "\n"
            
            # Add data rows
            for row in clean_table[1:]:
                table_text += " | ".join(row) + "\n"
        
        return table_text.strip()
    
    async def parse_excel(self, file_bytes: bytes, filename: str, include_raw_data: bool = False) -> Dict[str, Any]:
        """
        Parse Excel document and extract data from all sheets.
        
        Args:
            file_bytes: Excel file content as bytes
            filename: Original filename
            include_raw_data: If True, include raw headers and rows as lists
            
        Returns:
            Dictionary containing extracted data from all sheets
        """
        try:
            # Create file-like object from bytes
            excel_file = io.BytesIO(file_bytes)
            
            # Determine file type
            file_extension = Path(filename).suffix.lower()
            
            # Read Excel file with pandas
            if file_extension == '.xlsx':
                # Read all sheets
                excel_data = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
            elif file_extension == '.xls':
                # Read all sheets (older Excel format)
                excel_data = pd.read_excel(excel_file, sheet_name=None, engine='xlrd')
            else:
                raise ValueError(f"Unsupported Excel format: {file_extension}")
            
            # Initialize results
            extracted_sheets = []
            all_text_parts = []
            
            # Process each sheet
            for sheet_name, df in excel_data.items():
                # Skip empty sheets
                if df.empty:
                    continue
                
                # Convert DataFrame to text representation
                sheet_text = self._dataframe_to_text(df, sheet_name)
                all_text_parts.append(sheet_text)
                
                # Get sheet statistics - convert pandas/numpy types to Python native types
                non_empty_cells = int(df.count().sum())
                total_cells = int(df.shape[0] * df.shape[1])
                
                # Convert sample data and handle NaN values
                sample_data = []
                if len(df) > 0:
                    sample_df = df.head(3).fillna("")  # Replace NaN with empty string
                    sample_data = sample_df.to_dict('records')
                    # Convert any remaining numpy types to Python types
                    for record in sample_data:
                        for key, value in record.items():
                            if pd.isna(value):
                                record[key] = ""
                            elif hasattr(value, 'item'):  # numpy scalar
                                record[key] = value.item()
                
                # Extract raw data if requested
                raw_headers = None
                raw_rows = None
                
                if include_raw_data:
                    raw_headers = [str(col) for col in df.columns.tolist()]
                    # Convert DataFrame to list of lists, replacing NaN with empty string
                    raw_rows = df.fillna("").values.tolist()
                
                # Extract sheet data
                sheet_info = {
                    "sheet_name": sheet_name,
                    "text": sheet_text,
                    "rows": int(len(df)),
                    "columns": int(len(df.columns)),
                    "non_empty_cells": non_empty_cells,
                    "total_cells": total_cells,
                    "column_names": [str(col) for col in df.columns.tolist()],
                    "data_types": {str(k): str(v) for k, v in df.dtypes.to_dict().items()},
                    "sample_data": sample_data
                }
                
                # Add raw data if requested
                if include_raw_data:
                    sheet_info["raw_headers"] = raw_headers
                    sheet_info["raw_rows"] = raw_rows
                
                extracted_sheets.append(sheet_info)
            
            # Combine all sheet text
            full_text = "\n\n".join(all_text_parts)
            
            # Determine document type
            doc_type = DocumentType.XLSX if file_extension == '.xlsx' else DocumentType.XLS
            
            return {
                "document_type": doc_type,
                "filename": filename,
                "extracted_text": full_text,
                "sheets": extracted_sheets,
                "metadata": {
                    "parser": "pandas/openpyxl",
                    "total_characters": int(len(full_text)),
                    "sheet_count": int(len(extracted_sheets)),
                    "total_rows": int(sum(sheet["rows"] for sheet in extracted_sheets)),
                    "total_columns": int(sum(sheet["columns"] for sheet in extracted_sheets)),
                    "file_format": file_extension
                }
            }
            
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to parse Excel '{filename}': {str(e)}"
            )
    
    def _dataframe_to_text(self, df: pd.DataFrame, sheet_name: str) -> str:
        """
        Convert pandas DataFrame to readable text format.
        
        Args:
            df: Pandas DataFrame
            sheet_name: Name of the Excel sheet
            
        Returns:
            Formatted text representation of the data
        """
        if df.empty:
            return f"[Sheet: {sheet_name}]\n(Empty sheet)"
        
        # Start with sheet header
        text_parts = [f"[Sheet: {sheet_name}]"]
        
        # Add column headers
        headers = " | ".join(str(col) for col in df.columns)
        text_parts.append(headers)
        text_parts.append("-" * len(headers))
        
        # Add data rows (limit to prevent excessive text)
        max_rows = 100  # Limit rows to prevent huge text output
        for idx, row in df.head(max_rows).iterrows():
            row_text = " | ".join(str(val) if pd.notna(val) else "" for val in row)
            text_parts.append(row_text)
        
        # Add truncation notice if needed
        if len(df) > max_rows:
            text_parts.append(f"... ({len(df) - max_rows} more rows)")
        
        return "\n".join(text_parts)
    
    async def get_document_info(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Get basic document information without full parsing.
        
        Args:
            file_bytes: Document file content as bytes
            filename: Original filename
            
        Returns:
            Basic document metadata
        """
        try:
            # Determine document type from filename
            file_extension = Path(filename).suffix.lower()
            
            if file_extension == '.pdf':
                # Quick PDF info
                pdf_file = io.BytesIO(file_bytes)
                with pdfplumber.open(pdf_file) as pdf:
                    return {
                        "document_type": DocumentType.PDF,
                        "filename": filename,
                        "file_size": len(file_bytes),
                        "total_pages": len(pdf.pages),
                        "file_extension": file_extension
                    }
            elif file_extension == '.docx':
                # Quick DOCX info
                docx_file = io.BytesIO(file_bytes)
                doc = Document(docx_file)
                return {
                    "document_type": DocumentType.DOCX,
                    "filename": filename,
                    "file_size": len(file_bytes),
                    "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
                    "table_count": len(doc.tables),
                    "file_extension": file_extension
                }
            elif file_extension in ['.xlsx', '.xls']:
                # Quick Excel info
                excel_file = io.BytesIO(file_bytes)
                if file_extension == '.xlsx':
                    excel_data = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
                else:
                    excel_data = pd.read_excel(excel_file, sheet_name=None, engine='xlrd')
                
                doc_type = DocumentType.XLSX if file_extension == '.xlsx' else DocumentType.XLS
                total_rows = int(sum(len(df) for df in excel_data.values()))
                total_sheets = int(len(excel_data))
                
                return {
                    "document_type": doc_type,
                    "filename": filename,
                    "file_size": int(len(file_bytes)),
                    "sheet_count": total_sheets,
                    "total_rows": total_rows,
                    "sheet_names": [str(name) for name in excel_data.keys()],
                    "file_extension": file_extension
                }
            else:
                return {
                    "document_type": DocumentType.OTHER,
                    "filename": filename,
                    "file_size": len(file_bytes),
                    "file_extension": file_extension,
                    "supported": False
                }
                
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Failed to get document info for '{filename}': {str(e)}"
            )

# Global parser service instance
parser_service = ParserService()
